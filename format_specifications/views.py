from django.shortcuts import render
from django.http import FileResponse, HttpResponseBadRequest, HttpResponse, JsonResponse
from django.views.decorators.http import require_http_methods
from .utils import generate_output_path
from .utils.word_formatter import AIWordFormatter
from .utils.ai_word_utils import AITextProcessor
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime
from django.conf import settings
import logging
import tempfile
import re

# 获取logger实例
logger = logging.getLogger(__name__)

# Processing log helper
def add_processing_log(request, message):
    """Add a log entry to processing status session"""
    if 'processing' not in request.session:
        request.session['processing'] = {
            'status': 'processing',
            'logs': [],
            'current_step': 0,
            'total_steps': 0
        }

    log_entry = {
        'msg': message,
        'time': datetime.now().strftime('%H:%M:%S')
    }

    # Keep only last 50 log entries to prevent session bloat
    logs = request.session['processing']['logs']
    logs.append(log_entry)
    if len(logs) > 50:
        logs.pop(0)

    request.session['processing']['logs'] = logs
    request.session.modified = True  # Ensure session is saved
    logger.debug(f"Processing log: {message}")

# 上传页面（新增 AI 开关选项）
def upload_word_page(request):
    logger.info("访问上传页面")

    # Get templates for template generation section
    from .services.template_manager import TemplateManager
    templates = TemplateManager.list_available_templates(
        request.user if request.user.is_authenticated else None
    )

    logger.info(f"加载了 {len(templates)} 个模板")
    for template_id, name, category, template_type in templates[:3]:
        logger.info(f"  - 模板: {template_id} = {name} ({category})")

    return render(request, 'upload_word_ai.html', {
        'templates': templates
    })

# 处理 AI 辅助格式化
@require_http_methods(["POST"])
def ai_format_word(request):
    logger.info("开始处理AI格式化请求")

    # 1. Get optimization mode
    optimization_mode = request.POST.get('optimization_mode', 'simple')
    logger.info(f"Optimization mode: {optimization_mode}")

    # 2. Check file upload (always required)
    if 'word_file' not in request.FILES:
        error_msg = "请上传 Word 文件"
        logger.warning(error_msg)
        return render(request, 'upload_word_ai.html', {'error': error_msg})

    uploaded_file = request.FILES['word_file']
    if not uploaded_file.name.endswith(('.docx',)):
        error_msg = "仅支持 .docx 格式（.doc 需先转换为 .docx）"
        logger.warning(error_msg)
        return render(request, 'upload_word_ai.html', {'error': error_msg})

    # 3. Route based on mode
    if optimization_mode == 'simple':
        return handle_simple_optimization(request, uploaded_file)
    elif optimization_mode == 'template':
        return handle_template_optimization(request, uploaded_file)
    elif optimization_mode == 'custom':
        return handle_custom_optimization(request, uploaded_file)
    else:
        error_msg = f"无效的优化模式: {optimization_mode}"
        logger.error(error_msg)
        return render(request, 'upload_word_ai.html', {'error': error_msg})


def handle_simple_optimization(request, uploaded_file):
    """Handle simple optimization mode (existing functionality)"""
    logger.info("Processing simple optimization mode")

    # Check if AI is enabled
    use_ai_raw = request.POST.get('use_ai')
    use_ai = use_ai_raw == 'on'
    logger.info(f"AI功能启用状态: {use_ai} (原始值: '{use_ai_raw}')")

    # Get tone parameter
    tone = request.POST.get('tone', 'no_preference')
    logger.info(f"选择的语调: {tone}")

    # Get style template
    style_template = request.POST.get('style_template', 'default')

    # 调试：记录接收到的原始字号值
    raw_heading_size = request.POST.get('heading_size')
    raw_body_size = request.POST.get('body_size')
    logger.info(f"📥 接收到标题字号: '{raw_heading_size}' (类型: {type(raw_heading_size).__name__ if raw_heading_size else 'None'})")
    logger.info(f"📥 接收到正文字号: '{raw_body_size}' (类型: {type(raw_body_size).__name__ if raw_body_size else 'None'})")

    # 辅助函数：安全转换数值，空字符串返回 None
    def safe_int(value):
        if value and value.strip():
            try:
                return int(value)
            except (ValueError, TypeError):
                return None
        return None

    def safe_float(value):
        if value and value.strip():
            try:
                return float(value)
            except (ValueError, TypeError):
                return None
        return None

    # 处理自定义图片尺寸（支持英寸和厘米）
    image_width_value = request.POST.get('image_width')
    if image_width_value == 'custom':
        width_custom = safe_float(request.POST.get('image_width_custom'))
        width_unit = request.POST.get('image_width_unit', 'inch')
        if width_custom:
            # 如果是厘米，转换为英寸（1英寸 = 2.54厘米）
            image_width = width_custom if width_unit == 'inch' else width_custom / 2.54
        else:
            image_width = None
    else:
        image_width = safe_float(image_width_value)

    image_height_value = request.POST.get('image_height')
    if image_height_value == 'custom':
        height_custom = safe_float(request.POST.get('image_height_custom'))
        height_unit = request.POST.get('image_height_unit', 'inch')
        if height_custom:
            # 如果是厘米，转换为英寸（1英寸 = 2.54厘米）
            image_height = height_custom if height_unit == 'inch' else height_custom / 2.54
        else:
            image_height = None
    else:
        image_height = safe_float(image_height_value)

    custom_config = {
        'heading_font': request.POST.get('heading_font') or None,
        # 字号支持中文代码（如 'xiaosi'）或数值，不使用 safe_int 转换
        'heading_size': request.POST.get('heading_size') or None,
        'body_font': request.POST.get('body_font') or None,
        # 字号支持中文代码（如 'xiaosi'）或数值，不使用 safe_int 转换
        'body_size': request.POST.get('body_size') or None,
        'line_spacing': safe_float(request.POST.get('line_spacing')),
        'image_width': image_width,
        'image_height': image_height,
    }

    # 合并模板和自定义值
    from format_specifications.utils.word_formatter import STYLE_TEMPLATES
    style_config = STYLE_TEMPLATES[style_template].copy()
    for key, value in custom_config.items():
        if value is not None:
            style_config[key] = value

    logger.info(f"样式模板: {style_template}, 合并后配置: {style_config}")

    # 3. 保存上传文件
    upload_dir = os.path.join(settings.MEDIA_ROOT, 'uploaded_words')
    os.makedirs(upload_dir, exist_ok=True)
    input_file_path = os.path.join(upload_dir, uploaded_file.name)
    with open(input_file_path, 'wb') as f:
        for chunk in uploaded_file.chunks():
            f.write(chunk)

    # 4. 生成输出文件路径
    output_file_path, output_filename = generate_output_path(uploaded_file)
    print(output_filename)

    # 5. 执行 AI 格式化
    try:
        logger.info(f"开始格式化文件: {input_file_path}, 输出到: {output_file_path}")

        # Create log callback to send detailed AI logs to frontend
        def log_callback(message):
            """Callback to send AI processing logs to progress display"""
            add_processing_log(request, message)

        # 读取重试配置用于会话跟踪
        retry_enabled = getattr(settings, 'ZHIPU_RETRY_ENABLED', True)
        retry_count = getattr(settings, 'ZHIPU_RETRY_COUNT', 1)

        # 初始化会话状态
        request.session['ai_processing'] = {
            'status': 'processing',
            'attempt': 1,
            'max_attempts': retry_count + 1 if use_ai and retry_enabled else 1,
            'timestamp': datetime.now().isoformat()
        }
        request.session.save()

        formatter = AIWordFormatter(input_file_path, use_ai=use_ai, tone=tone, style_config=style_config, log_callback=log_callback)

        # 在格式化前获取原始文档分析数据
        original_analysis = formatter.analyze_document()
        logger.info(f"原始文档分析: {original_analysis}")

        result = formatter.format(output_file_path)
        logger.info("文件格式化完成")

        # 更新会话状态为成功
        request.session['ai_processing']['status'] = 'complete'
        request.session['ai_processing']['completed_at'] = datetime.now().isoformat()
        request.session.save()

        # 检查生成的文件是否为空
        if os.path.getsize(output_file_path) == 0:
            os.remove(output_file_path)
            raise ValueError("生成的文件为空，请重试")

        # 记录原始文件名和生成的文件名
        logger.info(f"原始文件名: {uploaded_file.name}, 生成文件名: {output_filename}")

        # 返回文件下载，设置正确的Content-Disposition头
        response = FileResponse(open(output_file_path, 'rb'))
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        # 使用引号包围文件名，确保浏览器正确处理包含中文的文件名

        from urllib.parse import quote  # 导入URL编码模块
        # 1. 对文件名做URL编码（解决中文/特殊字符）
        encoded_filename = quote(output_filename)
        response['Content-Disposition'] = (
            f'attachment; filename="{encoded_filename}"; '
            f'filename*=UTF-8\'\'{encoded_filename}'
        )

        return response

    except ValueError as ve:
        # AI返回空或文件为空的情况
        # 更新会话状态为失败
        request.session['ai_processing'] = {
            'status': 'failed',
            'error': str(ve),
            'timestamp': datetime.now().isoformat()
        }
        request.session.save()

        # 重新创建formatter以获取分析数据
        formatter = AIWordFormatter(input_file_path, use_ai=use_ai, tone=tone, style_config=style_config)
        original_analysis = formatter.analyze_document()
        return render(request, 'upload_word_ai.html', {
            'error': str(ve),
            'original_analysis': original_analysis
        })
    except Exception as e:
        # 其他错误
        # 更新会话状态为失败
        request.session['ai_processing'] = {
            'status': 'failed',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }
        request.session.save()

        # 尝试创建formatter以获取分析数据，即使失败也要显示
        try:
            formatter = AIWordFormatter(input_file_path, use_ai=use_ai, tone=tone, style_config=style_config)
            original_analysis = formatter.analyze_document()
        except:
            original_analysis = None
        return render(request, 'upload_word_ai.html', {
            'error': f"处理失败：{str(e)}",
            'original_analysis': original_analysis
        })


def handle_template_optimization(request, uploaded_file):
    """
    Handle template-based optimization mode
    """
    from .services.template_manager import TemplateManager
    from .utils.ai_word_utils import AITextProcessor
    from .utils.image_tracker import DocumentImageTracker, ImageReinsertionStrategy
    from docx import Document
    import re

    logger.info("Processing template optimization mode")

    # Initialize processing logs
    add_processing_log(request, "开始模板优化处理 / Starting template optimization")
    add_processing_log(request, f"正在处理文档: {uploaded_file.name}")

    # Get template selection
    template_id = request.POST.get('template_id')
    if not template_id:
        error_msg = "请选择模板 / Please select a template"
        logger.warning(error_msg)
        return render(request, 'upload_word_ai.html', {'error': error_msg})

    # Get template from TemplateManager
    template = TemplateManager.get_template(
        template_id,
        request.user if request.user.is_authenticated else None
    )

    if not template:
        error_msg = f"模板 '{template_id}' 不存在"
        logger.warning(error_msg)
        add_processing_log(request, f"❌ 错误: {error_msg}")
        request.session['processing']['status'] = 'failed'
        return render(request, 'upload_word_ai.html', {'error': error_msg})

    add_processing_log(request, f"加载模板: {template.name}")

    # Get common settings
    tone = request.POST.get('tone', 'no_preference')
    style_template = request.POST.get('style_template', 'default')

    # Get style config (reused from simple mode logic)
    from format_specifications.utils.word_formatter import STYLE_TEMPLATES
    style_config = STYLE_TEMPLATES[style_template].copy()

    # Save uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        for chunk in uploaded_file.chunks():
            tmp_file.write(chunk)
        tmp_file_path = tmp_file.name

    # Initialize image tracker
    image_tracker = DocumentImageTracker(tmp_file_path)
    extracted_images = []

    try:
        # Extract images with context BEFORE processing text
        try:
            extracted_images = image_tracker.extract_images_with_context()
            add_processing_log(request, f"检测到 {len(extracted_images)} 张图片 / Detected {len(extracted_images)} image(s)")
            logger.info(f"Extracted {len(extracted_images)} images from document")
        except Exception as e:
            logger.warning(f"Image extraction failed: {e}")
            add_processing_log(request, f"⚠️ 图片提取失败 / Image extraction failed: {str(e)}")

        # Extract text from uploaded document
        logger.info(f"Extracting text from document: {tmp_file_path}")
        add_processing_log(request, "提取文档内容... / Extracting document content")
        doc = Document(tmp_file_path)
        source_document_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        add_processing_log(request, f"提取完成: {len(source_document_text)} 字符 / Extracted {len(source_document_text)} chars")

        if not source_document_text.strip():
            error_msg = "上传的文档内容为空，无法处理"
            logger.warning(error_msg)
            return render(request, 'upload_word_ai.html', {'error': error_msg})

        logger.info(f"Extracted {len(source_document_text)} characters from source document")

        # Create log callback to send detailed AI logs to frontend
        def log_callback(message):
            """Callback to send AI processing logs to progress display"""
            add_processing_log(request, message)

        # Choose processing mode based on document size
        processor = AITextProcessor(tone=tone, log_callback=log_callback)

        if len(source_document_text) < 500:
            # Use batch mode for small documents
            logger.info(f"Small document ({len(source_document_text)} chars), using BATCH mode")
            add_processing_log(request, f"使用批量处理模式 / Using BATCH mode ({len(source_document_text)} chars)")
            generated_content = processor.generate_from_template_batch(
                template=template,
                source_document_text=source_document_text,
                user_outline="",
                tone=tone
            )
        else:
            # Use sequential mode for medium/large documents
            logger.info(f"Medium/Large document ({len(source_document_text)} chars), using SEQUENTIAL mode")
            add_processing_log(request, f"使用顺序处理模式 / Using SEQUENTIAL mode ({len(source_document_text)} chars)")
            add_processing_log(request, f"处理 {len(template.sections)} 个章节 / Processing {len(template.sections)} sections")
            generated_content = processor.generate_from_template(
                template=template,
                user_outline="",  # No user outline, use extracted content
                source_document_text=source_document_text,
                tone=tone
            )

        # Match extracted images to sections based on context
        image_insertions = []
        if extracted_images:
            add_processing_log(request, f"匹配图片到章节 / Matching images to sections ({len(extracted_images)} images)")
            for image_meta in extracted_images:
                section_id, position = ImageReinsertionStrategy.find_best_insertion_position(
                    image_meta,
                    generated_content,
                    template
                )
                image_insertions.append({
                    'section_id': section_id,
                    'image_path': image_meta['image_path'],
                    'position': position
                })
                logger.debug(f"Matched image to section: {section_id}")

            add_processing_log(request, f"已匹配 {len(image_insertions)} 张图片 / Matched {len(image_insertions)} image(s)")

        # Build document from generated content
        output_file_path, output_filename = generate_output_path(uploaded_file)
        output_doc = Document()

        # Get style config for image dimensions
        from docx.shared import Inches
        image_width = Inches(style_config['image_width'])
        image_height = Inches(style_config['image_height'])

        # Add title
        title = output_doc.add_heading(template.name, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Helper function to check if content is meaningful (not placeholder or too short)
        def is_meaningful_content(content: str) -> bool:
            """Check if content has meaningful text (not just placeholders)"""
            if not content or not isinstance(content, str):
                return False

            content = content.strip()

            # Check if too short (less than 15 characters)
            if len(content) < 15:
                return False

            # Check for placeholder patterns
            placeholders = [
                '[待补充]', '[待填写]', '[待完善]',
                '待补充', '待填写', '待完善',
                '请补充', '请填写', '请完善'
            ]

            # If content is mostly placeholders, it's not meaningful
            placeholder_count = sum(1 for p in placeholders if p in content)
            if placeholder_count > 0:
                # Calculate what percentage of content is placeholders
                total_placeholder_chars = sum(len(p) for p in placeholders if p in content)
                if total_placeholder_chars / len(content) > 0.3:  # More than 30% placeholders
                    return False

            return True

        # Debug logging
        logger.info(f"Filtering and writing sections (only meaningful content)")
        sections_written = 0
        sections_skipped = 0

        # Add sections based on template structure
        for section in template.sections:
            # Check if main section has meaningful content
            section_has_content = False
            section_content = None

            if section.id in generated_content:
                section_content = generated_content[section.id]
                if is_meaningful_content(section_content):
                    section_has_content = True

            # Check if any subsection has meaningful content
            subsections_with_content = []
            for subsection in section.subsections:
                if subsection.id in generated_content:
                    subsection_content = generated_content[subsection.id]
                    if is_meaningful_content(subsection_content):
                        subsections_with_content.append((subsection, subsection_content))

            # Only write section if it has content OR has subsections with content
            if section_has_content or subsections_with_content:
                # Write main section heading and content (if exists)
                if section_has_content:
                    output_doc.add_heading(section.title, 1)
                    # ✅ 修复：添加段落样式设置，包括首行缩进
                    content_para = output_doc.add_paragraph(section_content)
                    from docx.oxml.ns import qn
                    content_para.paragraph_format.line_spacing = 1.5
                    content_para.paragraph_format.first_line_indent = Pt(21.0)
                    # Set font properties
                    for run in content_para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(12)
                    logger.info(f"✓ Wrote section: {section.title} ({len(section_content)} chars)")
                    sections_written += 1

                    # Insert images matched to this section
                    section_images = [img for img in image_insertions if img['section_id'] == section.id]
                    for img_data in section_images:
                        from docx.shared import Pt
                        img_para = output_doc.add_paragraph()
                        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        img_para.paragraph_format.space_after = Pt(12)
                        img_para.paragraph_format.space_before = Pt(12)

                        try:
                            img_run = img_para.add_run()
                            img_run.add_picture(
                                img_data['image_path'],
                                width=image_width,
                                height=image_height
                            )
                            logger.info(f"  ✓ Inserted image in section: {section.title}")
                        except Exception as e:
                            logger.warning(f"  ✗ Failed to insert image: {e}")
                            img_para.add_run("[图片加载失败 / Image load failed]")
                elif subsections_with_content:
                    # Section has no main content but has subsections - write heading only
                    output_doc.add_heading(section.title, 1)
                    logger.info(f"✓ Wrote section heading (no main content): {section.title}")
                    sections_written += 1

                # Write subsections with meaningful content
                for subsection, subsection_content in subsections_with_content:
                    output_doc.add_heading(subsection.title, 2)
                    # ✅ 修复：添加段落样式设置，包括首行缩进
                    subsection_para = output_doc.add_paragraph(subsection_content)
                    from docx.oxml.ns import qn
                    subsection_para.paragraph_format.line_spacing = 1.5
                    subsection_para.paragraph_format.first_line_indent = Pt(21.0)
                    # Set font properties
                    for run in subsection_para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(12)
                    logger.info(f"  ✓ Wrote subsection: {subsection.title} ({len(subsection_content)} chars)")
            else:
                # Skip this section entirely - no title, no content
                logger.info(f"⊘ Skipping section (no meaningful content): {section.title}")
                sections_skipped += 1

        logger.info(f"Document generation complete: {sections_written} sections written, {sections_skipped} sections skipped")

        # Save document
        add_processing_log(request, f"✅ 生成文档完成 / Document generated ({sections_written} sections written)")
        request.session['processing']['status'] = 'complete'
        output_doc.save(output_file_path)
        logger.info(f"Generated document saved to: {output_file_path}")

        # Return file response
        response = FileResponse(open(output_file_path, 'rb'))
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        from urllib.parse import quote
        encoded_filename = quote(output_filename)
        response['Content-Disposition'] = (
            f'attachment; filename="{encoded_filename}"; '
            f'filename*=UTF-8\'\'{encoded_filename}'
        )

        return response

    except Exception as e:
        logger.error(f"Template optimization failed: {str(e)}")
        error_msg = f"模板优化失败: {str(e)}"
        return render(request, 'upload_word_ai.html', {'error': error_msg})

    finally:
        # Clean up temporary file
        if os.path.exists(tmp_file_path):
            os.unlink(tmp_file_path)

        # Clean up extracted images
        image_tracker.cleanup()


def find_best_custom_section_for_image(image_metadata: dict, structure_sections: list) -> str:
    """
    Find the best section title to insert an image for custom structure mode.

    Uses a multi-strategy approach:
    1. Primary: Keyword relevance scoring against section titles
    2. Fallback 1: First section with substantial content (>100 chars)
    3. Fallback 2: Last section
    4. Last resort: First section

    Args:
        image_metadata: Image context from original document
        structure_sections: List of section dicts with 'title' keys

    Returns:
        Section title to insert image into (None if no match)
    """
    from typing import List, Dict

    best_section = None
    best_score = 0.0

    # Try to match based on keywords in section titles
    for section in structure_sections:
        section_title = section['title']

        # Calculate relevance score
        score = 0.0

        # Check section title against preceding text
        if section_title.lower() in image_metadata['preceding_text'].lower():
            score += 0.5

        # Check section title against following text
        if section_title.lower() in image_metadata['following_text'].lower():
            score += 0.5

        # Check section title against paragraph text
        if section_title.lower() in image_metadata['paragraph_text'].lower():
            score += 0.3

        if score > best_score:
            best_score = score
            best_section = section_title

    # If best match found and has meaningful score
    if best_section and best_score > 0:
        return best_section

    # Fallback: first section (if any sections exist)
    if structure_sections and len(structure_sections) > 0:
        return structure_sections[0]['title']

    return None


def handle_custom_optimization(request, uploaded_file):
    """
    Handle custom structure optimization mode
    """
    from .utils.ai_word_utils import AITextProcessor
    from .utils.image_tracker import DocumentImageTracker
    from docx import Document
    from docx.shared import Pt
    import re

    logger.info("Processing custom structure optimization mode")

    # Get custom structure points
    custom_structure = request.POST.get('custom_structure', '').strip()

    if not custom_structure:
        error_msg = "请输入自定义结构要点"
        logger.warning(error_msg)
        return render(request, 'upload_word_ai.html', {'error': error_msg})

    # Parse structure points
    structure_sections = parse_custom_structure(custom_structure)
    logger.info(f"Parsed {len(structure_sections)} structure sections")

    if not structure_sections:
        error_msg = "无法解析自定义结构，请确保每行一个结构要点"
        logger.warning(error_msg)
        return render(request, 'upload_word_ai.html', {'error': error_msg})

    # Get common settings
    tone = request.POST.get('tone', 'no_preference')
    style_template = request.POST.get('style_template', 'default')

    # Get style config
    from format_specifications.utils.word_formatter import STYLE_TEMPLATES
    style_config = STYLE_TEMPLATES[style_template].copy()

    # Save uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        for chunk in uploaded_file.chunks():
            tmp_file.write(chunk)
        tmp_file_path = tmp_file.name

    # Initialize image tracker
    image_tracker = DocumentImageTracker(tmp_file_path)
    extracted_images = []

    try:
        # Extract images with context BEFORE processing text
        try:
            extracted_images = image_tracker.extract_images_with_context()
            logger.info(f"Extracted {len(extracted_images)} images from document")
            add_processing_log(request, f"检测到 {len(extracted_images)} 张图片 / Detected {len(extracted_images)} image(s)")
        except Exception as e:
            logger.warning(f"Image extraction failed: {e}")
            add_processing_log(request, f"⚠️ 图片提取失败 / Image extraction failed: {str(e)}")

        # Extract text from uploaded document
        logger.info(f"Extracting text from document: {tmp_file_path}")
        doc = Document(tmp_file_path)
        source_document_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

        if not source_document_text.strip():
            error_msg = "上传的文档内容为空，无法处理"
            logger.warning(error_msg)
            return render(request, 'upload_word_ai.html', {'error': error_msg})

        # Create log callback to send detailed AI logs to frontend
        def log_callback(message):
            """Callback to send AI processing logs to progress display"""
            add_processing_log(request, message)

        # Generate content according to custom structure
        processor = AITextProcessor(tone=tone, log_callback=log_callback)
        generated_content = generate_with_custom_structure(
            processor,
            source_document_text,
            structure_sections
        )

        # Match extracted images to sections based on context
        image_insertions = []
        if extracted_images:
            add_processing_log(request, f"📷 匹配图片到章节 / Matching images to sections ({len(extracted_images)} images)")
            logger.info(f"Starting image matching for {len(extracted_images)} images to {len(structure_sections)} sections")

            for idx, image_meta in enumerate(extracted_images):
                section_title = find_best_custom_section_for_image(
                    image_meta,
                    structure_sections
                )
                if section_title:
                    image_insertions.append({
                        'section_title': section_title,
                        'image_path': image_meta['image_path']
                    })
                    logger.info(f"  Image {idx + 1}: matched to section '{section_title}' (context: '{image_meta['paragraph_text'][:50]}')")
                else:
                    logger.warning(f"  Image {idx + 1}: no matching section found (context: '{image_meta['paragraph_text'][:50]}')")

            logger.info(f"Image matching complete: {len(image_insertions)}/{len(extracted_images)} images matched")
            add_processing_log(request, f"✅ 已匹配 {len(image_insertions)} 张图片 / Matched {len(image_insertions)} image(s)")
        else:
            logger.warning("No extracted images to match")
            add_processing_log(request, "⚠️ 未检测到图片 / No images detected")

        # Build document
        output_file_path, output_filename = generate_output_path(uploaded_file)
        output_doc = Document()

        # Get style config for image dimensions
        from docx.shared import Inches
        image_width = Inches(style_config['image_width'])
        image_height = Inches(style_config['image_height'])
        logger.info(f"Image dimensions: {style_config['image_width']}\" x {style_config['image_height']}\"")

        # Add title
        title = output_doc.add_heading('自定义结构文档', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add sections based on custom structure
        sections_with_images = 0
        total_images_inserted = 0

        for section in structure_sections:
            section_title = section['title']
            # ✅ 修复：即使内容为空，也添加章节标题和图片
            if section_title in generated_content:
                section_content = generated_content[section_title]

                # Add section heading (always add if section exists)
                output_doc.add_heading(section_title, 1)

                # Add section content with proper formatting
                if section_content and section_content.strip():
                    content_para = output_doc.add_paragraph(section_content)
                    # ✅ 修复：应用段落样式，包括首行缩进2格（21磅）
                    from docx.oxml.ns import qn
                    content_para.paragraph_format.line_spacing = 1.5
                    content_para.paragraph_format.first_line_indent = Pt(21.0)

                    # Set font properties for content
                    for run in content_para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(12)

                # Insert images matched to this section (even if content is empty)
                section_images = [img for img in image_insertions if img['section_title'] == section_title]
                if section_images:
                    sections_with_images += 1
                    logger.info(f"Section '{section_title}': inserting {len(section_images)} image(s)")

                for img_idx, img_data in enumerate(section_images):
                    img_para = output_doc.add_paragraph()
                    img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    img_para.paragraph_format.space_after = Pt(12)
                    img_para.paragraph_format.space_before = Pt(12)

                    try:
                        img_run = img_para.add_run()
                        img_run.add_picture(
                            img_data['image_path'],
                            width=image_width,
                            height=image_height
                        )
                        total_images_inserted += 1
                        logger.info(f"  ✓ Inserted image {img_idx + 1} in section: {section_title}")
                        add_processing_log(request, f"  ✓ 插入图片到: {section_title} / Inserted image in: {section_title}")
                    except Exception as e:
                        logger.warning(f"  ✗ Failed to insert image: {e}")
                        img_para.add_run("[图片加载失败 / Image load failed]")

        # Save document
        logger.info(f"Document generation complete:")
        logger.info(f"  - Sections with images: {sections_with_images}/{len(structure_sections)}")
        logger.info(f"  - Total images inserted: {total_images_inserted}/{len(extracted_images)}")

        output_doc.save(output_file_path)
        logger.info(f"Generated document saved to: {output_file_path}")

        # Return file response
        response = FileResponse(open(output_file_path, 'rb'))
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        from urllib.parse import quote
        encoded_filename = quote(output_filename)
        response['Content-Disposition'] = (
            f'attachment; filename="{encoded_filename}"; '
            f'filename*=UTF-8\'\'{encoded_filename}'
        )

        return response

    except Exception as e:
        logger.error(f"Custom structure optimization failed: {str(e)}")
        error_msg = f"自定义结构优化失败: {str(e)}"
        return render(request, 'upload_word_ai.html', {'error': error_msg})

    finally:
        # Clean up temporary file
        if os.path.exists(tmp_file_path):
            os.unlink(tmp_file_path)

        # Clean up extracted images
        image_tracker.cleanup()


def parse_custom_structure(structure_text):
    """
    Parse user's custom structure into sections
    Supports:
    - Numbered list: 1. Title, 2. Title
    - Plain lines: Just text (treat as section title)
    """
    lines = structure_text.strip().split('\n')
    sections = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Remove numbering if present
        clean_title = line
        if line[0].isdigit() and ('.' in line or ')' in line or '、' in line):
            # Extract title after number
            parts = re.split(r'^[\d]+[\.\）、]\s*', line, 1)
            if len(parts) > 1:
                clean_title = parts[1]

        sections.append({
            'title': clean_title,
            'original': line
        })

    return sections


def generate_with_custom_structure(processor, source_text, structure_sections):
    """
    Generate document content organized by custom structure
    """
    generated_content = {}

    for section in structure_sections:
        # Extract content relevant to this section
        extracted = processor.extract_section_for_structure(
            source_text,
            section['title']
        )

        # Fallback: if extraction failed, use first 1000 chars of source text
        if not extracted or not extracted.strip():
            logger.warning(f"Extraction failed for section '{section['title']}', using fallback")
            extracted = source_text[:1000]

        logger.info(f"Extracted {len(extracted) if extracted else 0} chars for section: {section['title']}")

        # Polish the extracted content
        if extracted and extracted.strip() and len(extracted.strip()) > 10:
            logger.info(f"Polishing content for section: {section['title']}")
            polished = processor.process_text(extracted)
            generated_content[section['title']] = polished
            logger.info(f"Polished content: {len(polished)} chars")
        else:
            logger.warning(f"Section {section['title']} has no meaningful content, skipping polishing")
            generated_content[section['title']] = ""

    return generated_content


@require_http_methods(["GET"])
def ai_processing_status(request):
    """
    返回 AI 处理状态，用于前端轮询

    性能优化：
    - 不调用 AI API，仅读取 session
    - 响应时间 < 10ms
    - 前端轮询间隔 >= 4 秒
    """
    # 快速返回 session 中的状态（不涉及任何 AI 调用）
    processing_info = request.session.get('ai_processing', {
        'status': 'unknown'
    })

    # 添加响应头，防止浏览器缓存
    response = JsonResponse(processing_info)
    response['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'

    return response


@require_http_methods(["GET"])
def processing_status(request):
    """
    返回通用处理状态，用于前端轮询（支持所有优化模式）

    性能优化：
    - 不调用 AI API，仅读取 session
    - 响应时间 < 10ms
    - 前端轮询间隔 2 秒
    """
    # 快速返回 session 中的处理状态（不涉及任何 AI 调用）
    processing_info = request.session.get('processing', {
        'status': 'unknown',
        'logs': [],
        'current_step': 0,
        'total_steps': 0
    })

    # 添加响应头，防止浏览器缓存
    response = JsonResponse(processing_info)
    response['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'

    return response


# ==================== 文档分割功能 (Document Segmentation) ====================

def segmentation_only_page(request):
    """
    渲染文档分割页面
    """
    logger.info("访问文档分割页面")
    return render(request, 'segmentation_only.html')


@require_http_methods(["POST"])
def segment_document(request):
    """
    处理文档分割请求

    参数:
    - document: 上传的Word文档
    - mode: 分割模式 (paragraph/sentence/semantic)
    - include_metadata: 是否包含元数据 (可选)

    返回:
    - 分割后的Word文档下载
    """
    logger.info("开始处理文档分割请求")

    # 1. 验证文件上传
    if 'document' not in request.FILES:
        error_msg = "请上传Word文档"
        logger.warning(error_msg)
        return render(request, 'segmentation_only.html', {'error': error_msg})

    uploaded_file = request.FILES['document']
    if not uploaded_file.name.endswith('.docx'):
        error_msg = "仅支持 .docx 格式"
        logger.warning(error_msg)
        return render(request, 'segmentation_only.html', {'error': error_msg})

    # 2. 获取分割参数
    mode = request.POST.get('mode', 'paragraph')
    include_metadata = request.POST.get('include_metadata') == 'on'

    # 验证分割模式
    valid_modes = ['paragraph', 'sentence', 'semantic']
    if mode not in valid_modes:
        error_msg = f"无效的分割模式: {mode}"
        logger.warning(error_msg)
        return render(request, 'segmentation_only.html', {'error': error_msg})

    logger.info(f"分割模式: {mode}, 包含元数据: {include_metadata}")

    # 3. 保存上传的文件到临时位置
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            for chunk in uploaded_file.chunks():
                tmp_file.write(chunk)
            tmp_file_path = tmp_file.name

        logger.info(f"文件已保存到临时位置: {tmp_file_path}")

        # 4. 提取文档文本
        doc = Document(tmp_file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 跳过空段落
                paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)
        logger.info(f"提取了 {len(paragraphs)} 个段落，共 {len(full_text)} 个字符")

        # 5. 调用分割方法
        processor = AITextProcessor()

        if include_metadata:
            segments = processor.segment_text(full_text, mode=mode, include_metadata=True)
            logger.info(f"分割完成，共 {len(segments)} 个片段（包含元数据）")
        else:
            segments = processor.segment_text(full_text, mode=mode, include_metadata=False)
            logger.info(f"分割完成，共 {len(segments)} 个片段")

        # 6. 构建输出文档
        output_filename = f"segmented_{mode}_{uploaded_file.name}"

        # Create output directory
        output_dir = os.path.join(settings.MEDIA_ROOT, 'segmented')
        os.makedirs(output_dir, exist_ok=True)

        # Generate output path
        output_path = os.path.join(output_dir, output_filename)

        _build_segmented_document(segments, mode, include_metadata, output_path)
        logger.info(f"分割文档已保存到: {output_path}")

        # 7. 清理临时文件
        try:
            os.unlink(tmp_file_path)
        except:
            pass

        # 8. 返回文件
        response = FileResponse(open(output_path, 'rb'))
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response['Content-Disposition'] = f'attachment; filename="{output_filename}"'

        return response

    except Exception as e:
        logger.error(f"文档分割失败: {str(e)}", exc_info=True)
        error_msg = f"分割失败: {str(e)}"
        return render(request, 'segmentation_only.html', {'error': error_msg})


def _build_segmented_document(segments, mode, include_metadata, output_path):
    """
    构建分割后的Word文档

    参数:
    - segments: 分割后的片段列表或字典列表
    - mode: 分割模式
    - include_metadata: 是否包含元数据
    - output_path: 输出文件路径
    """
    doc = Document()

    # 添加标题
    title_text = {
        'paragraph': '文档分割结果 - 按段落分割',
        'sentence': '文档分割结果 - 按句子分割',
        'semantic': '文档分割结果 - 按语义分割'
    }

    title = doc.add_heading(title_text.get(mode, '文档分割结果'), 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加分割信息
    info_para = doc.add_paragraph()
    info_para.add_run(f"分割模式: {mode}\n")
    info_para.add_run(f"片段数量: {len(segments)}\n")
    info_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    info_para.runs[0].font.size = Pt(10)
    info_para.runs[0].font.color.rgb = None  # 使用默认灰色

    # 添加分隔线
    doc.add_paragraph("_" * 80)

    # 添加每个片段
    if include_metadata:
        # 包含元数据的格式
        for i, segment in enumerate(segments, 1):
            # 元数据信息
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"[片段 {i}]").bold = True
            meta_para.add_run(f" 类型: {segment.get('type', mode)} | ")
            meta_para.add_run(f"位置: {segment.get('position', i-1)}")
            meta_para.runs[0].font.size = Pt(9)
            meta_para.runs[0].font.color.rgb = None

            # 片段内容
            content_para = doc.add_paragraph(segment.get('text', ''))
            content_para.runs[0].font.size = Pt(11)

            # 片段间空行
            doc.add_paragraph()
    else:
        # 简单格式，只显示文本
        for i, segment in enumerate(segments, 1):
            if isinstance(segment, str):
                # 纯文本片段
                content_para = doc.add_paragraph(f"[{i}] {segment}")
                content_para.runs[0].font.size = Pt(11)
            else:
                # 字典形式（兼容性处理）
                content_para = doc.add_paragraph(f"[{i}] {segment.get('text', segment)}")
                content_para.runs[0].font.size = Pt(11)

    # 保存文档
    doc.save(output_path)
    logger.info(f"文档构建完成，共 {len(segments)} 个片段")


# ==================== 模板生成功能 (Template-Based Generation) ====================

def template_generation_page(request):
    """
    渲染模板生成页面
    """
    logger.info("访问模板生成页面")

    from format_specifications.services.template_manager import TemplateManager

    # Get all available templates
    user = request.user if request.user.is_authenticated else None
    templates = TemplateManager.list_available_templates(user)

    return render(request, 'template_generation.html', {
        'templates': templates
    })


@require_http_methods(["POST"])
def generate_from_template(request):
    """
    处理模板生成请求

    参数:
    - template_id: 选择的模板ID
    - user_outline: 用户提供的要点/大纲
    - source_document: 可选的源文档
    - tone: 可选的语调参数

    返回:
    - 生成的Word文档下载
    """
    logger.info("开始处理模板生成请求")

    from format_specifications.services.template_manager import TemplateManager
    from format_specifications.utils.document_extractor import DocumentExtractor
    from .utils.image_tracker import DocumentImageTracker, ImageReinsertionStrategy

    start_time = datetime.now()
    image_tracker = None
    tmp_file_path = None

    try:
        # 1. 获取模板ID
        template_id = request.POST.get('template_id')
        if not template_id:
            error_msg = "请选择一个模板"
            logger.warning(error_msg)
            return render(request, 'template_generation.html', {'error': error_msg})

        # 2. 获取用户要点
        user_outline = request.POST.get('user_outline', '').strip()
        if not user_outline:
            error_msg = "请提供文档要点"
            logger.warning(error_msg)
            return render(request, 'template_generation.html', {'error': error_msg})

        # 3. 获取可选的源文档
        source_document_text = None
        had_source_document = False
        extracted_images = []

        if 'source_document' in request.FILES and request.FILES['source_document']:
            uploaded_file = request.FILES['source_document']

            if uploaded_file.name.endswith('.docx'):
                # 保存临时文件
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    for chunk in uploaded_file.chunks():
                        tmp_file.write(chunk)
                    tmp_file_path = tmp_file.name

                # 提取文本
                try:
                    source_document_text = DocumentExtractor.extract_full_text(tmp_file_path)
                    had_source_document = True
                    logger.info(f"源文档已提取，共 {len(source_document_text)} 个字符")
                except Exception as e:
                    logger.warning(f"Failed to extract text from source document: {e}")

                # 提取图片
                try:
                    image_tracker = DocumentImageTracker(tmp_file_path)
                    extracted_images = image_tracker.extract_images_with_context()
                    logger.info(f"从源文档提取了 {len(extracted_images)} 张图片")
                except Exception as e:
                    logger.warning(f"图片提取失败: {str(e)}")

        # 4. 获取语调
        tone = request.POST.get('tone', 'no_preference')

        # 5. 获取模板
        user = request.user if request.user.is_authenticated else None
        template = TemplateManager.get_template(template_id, user)

        if not template:
            error_msg = f"模板 '{template_id}' 不存在"
            logger.warning(error_msg)
            return render(request, 'template_generation.html', {'error': error_msg})

        # 6. 生成内容
        logger.info(f"开始生成文档: 模板={template.name}, 语调={tone}")
        processor = AITextProcessor(tone=tone)

        generated_content = processor.generate_from_template(
            template=template,
            user_outline=user_outline,
            source_document_text=source_document_text,
            tone=tone
        )

        logger.info(f"内容生成完成，共 {len(generated_content)} 个章节")

        # 7. 获取样式配置 (用于图片尺寸)
        from format_specifications.utils.word_formatter import STYLE_TEMPLATES
        style_template = request.POST.get('style_template', 'default')
        style_config = STYLE_TEMPLATES[style_template].copy()

        # 8. 构建Word文档
        output_filename = f"{template.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # Create output directory
        output_dir = os.path.join(settings.MEDIA_ROOT, 'generated_from_template')
        os.makedirs(output_dir, exist_ok=True)

        # Generate output path
        output_path = os.path.join(output_dir, output_filename)

        _build_document_from_template(template, generated_content, output_path, extracted_images=extracted_images, style_config=style_config)
        logger.info(f"文档已生成: {output_path}")

        # 9. 计算耗时
        duration = int((datetime.now() - start_time).total_seconds())

        # 10. 记录使用日志
        try:
            TemplateManager.log_template_usage(
                template=template,
                user=user,
                user_outline=user_outline,
                had_source_document=had_source_document,
                generation_success=True,
                generation_duration=duration
            )
        except Exception as e:
            logger.warning(f"记录使用日志失败: {str(e)}")

        # 11. 返回文件
        response = FileResponse(open(output_path, 'rb'))
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        from urllib.parse import quote
        encoded_filename = quote(output_filename)
        response['Content-Disposition'] = (
            f'attachment; filename="{encoded_filename}"; '
            f'filename*=UTF-8\'\'{encoded_filename}'
        )

        return response

    except Exception as e:
        logger.error(f"模板生成失败: {str(e)}", exc_info=True)

        # 记录失败日志
        duration = int((datetime.now() - start_time).total_seconds())
        try:
            user = request.user if request.user.is_authenticated else None
            TemplateManager.log_template_usage(
                template=None,
                user=user,
                user_outline=request.POST.get('user_outline', ''),
                had_source_document='source_document' in request.FILES,
                generation_success=False,
                error_message=str(e),
                generation_duration=duration
            )
        except:
            pass

        error_msg = f"生成失败: {str(e)}"
        return render(request, 'template_generation.html', {'error': error_msg})

    finally:
        # Clean up temporary file
        if tmp_file_path and os.path.exists(tmp_file_path):
            try:
                os.unlink(tmp_file_path)
            except:
                pass

        # Clean up extracted images
        if image_tracker:
            try:
                image_tracker.cleanup()
            except:
                pass


@require_http_methods(["GET"])
def api_template_details(request, template_id):
    """
    AJAX端点：获取模板详细信息

    参数:
    - template_id: 模板ID

    返回:
    - JSON格式的模板详细信息
    """
    logger.info(f"获取模板详情: {template_id}")

    from format_specifications.services.template_manager import TemplateManager

    try:
        user = request.user if request.user.is_authenticated else None
        template_dict = TemplateManager.get_template_details_dict(template_id, user)

        if not template_dict:
            return JsonResponse({
                'success': False,
                'error': f"模板 '{template_id}' 不存在"
            }, status=404)

        return JsonResponse({
            'success': True,
            'template': template_dict
        })

    except Exception as e:
        logger.error(f"获取模板详情失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


def _build_document_from_template(template, generated_content, output_path, extracted_images=None, style_config=None):
    """
    根据模板和生成的内容构建Word文档

    参数:
    - template: 模板对象
    - generated_content: 字典，key为section_id，value为生成的内容
    - output_path: 输出文件路径
    - extracted_images: 提取的图片元数据列表 (可选)
    - style_config: 样式配置，用于图片尺寸 (可选)
    """
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from .utils.image_tracker import ImageReinsertionStrategy
    from docx.shared import Inches

    doc = Document()

    # 匹配图片到章节
    image_insertions = []
    if extracted_images:
        for image_meta in extracted_images:
            section_id, position = ImageReinsertionStrategy.find_best_insertion_position(
                image_meta,
                generated_content,
                template
            )
            if section_id:
                image_insertions.append({
                    'section_id': section_id,
                    'image_path': image_meta['image_path']
                })
                logger.debug(f"Matched image to section: {section_id}")

        logger.info(f"Matched {len(image_insertions)} images to sections")

    # 获取图片尺寸配置
    image_width = None
    image_height = None
    if style_config:
        try:
            image_width = Inches(style_config.get('image_width', 4.0))
            image_height = Inches(style_config.get('image_height', 3.0))
        except:
            logger.warning("Failed to get image dimensions from style config")

    # 添加标题
    title = doc.add_heading(template.name, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加生成时间
    info_para = doc.add_paragraph()
    info_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    info_para.add_run(f"模板: {template.name}")
    info_para.runs[0].font.size = Pt(10)
    info_para.runs[0].font.color.rgb = None

    # 添加分隔线
    doc.add_paragraph("_" * 80)

    # 遍历模板的章节结构
    def add_section(section, level=1):
        """递归添加章节及其内容"""

        # 获取生成的内容
        content = generated_content.get(section.id, '')

        # 如果有内容，添加到文档
        if content and content.strip():
            # 添加章节标题
            heading_level = min(level, 9)  # Word最多支持9级标题
            doc.add_heading(section.title, heading_level)

            # 添加内容
            # ✅ 修复：添加段落样式设置，包括首行缩进
            content_para = doc.add_paragraph(content)
            from docx.oxml.ns import qn
            content_para.paragraph_format.line_spacing = 1.5
            content_para.paragraph_format.first_line_indent = Pt(21.0)
            # Set font properties
            for run in content_para.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)

            # 插入匹配到该章节的图片
            section_images = [img for img in image_insertions if img['section_id'] == section.id]
            for img_data in section_images:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                img_para.paragraph_format.space_after = Pt(12)
                img_para.paragraph_format.space_before = Pt(12)

                try:
                    img_run = img_para.add_run()
                    img_run.add_picture(
                        img_data['image_path'],
                        width=image_width,
                        height=image_height
                    )
                    logger.info(f"  ✓ Inserted image in section: {section.title}")
                except Exception as e:
                    logger.warning(f"  ✗ Failed to insert image: {e}")
                    img_para.add_run("[图片加载失败 / Image load failed]")

        # 递归添加子章节
        for subsection in section.subsections:
            add_section(subsection, level=level + 1)

    # 从顶级章节开始
    for section in template.sections:
        if not section.is_optional or generated_content.get(section.id):
            add_section(section)

    # 保存文档
    doc.save(output_path)
    logger.info(f"文档构建完成")
