from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import logging
import os
import zipfile
import shutil
from pathlib import Path

# 导入 AI 处理类
from .ai_word_utils import AITextProcessor

# 配置日志
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

# 中文字号映射表（中文代号 -> 磅值）
CHINESE_FONT_SIZES = {
    # Pinyin keys (for backward compatibility)
    'chuhao': 42,      # 初号
    'xiaochu': 36,     # 小初
    'yihao': 26,       # 一号
    'xiaoyi': 24,      # 小一
    'erhao': 22,       # 二号
    'xiaoer': 18,      # 小二
    'sanhao': 16,      # 三号
    'xiaosan': 15,     # 小三
    'sihao': 14,       # 四号
    'xiaosi': 12,      # 小四
    'wuhao': 10.5,     # 五号
    'xiaowu': 9,       # 小五
    # Chinese character keys (for form submission)
    '初号': 42,
    '小初': 36,
    '一号': 26,
    '小一': 24,
    '二号': 22,
    '小二': 18,
    '三号': 16,
    '小三': 15,
    '四号': 14,
    '小四': 12,
    '五号': 10.5,
    '小五': 9,
}

# 反向映射：磅值 -> 中文名称（用于显示）
POINT_SIZE_TO_CHINESE = {
    42: '初号',
    36: '小初',
    26: '一号',
    24: '小一',
    22: '二号',
    18: '小二',
    16: '三号',
    15: '小三',
    14: '四号',
    12: '小四',
    10.5: '五号',
    9: '小五',
}

# 样式模板配置
STYLE_TEMPLATES = {
    'default': {
        'heading_font': '黑体',
        'heading_size': 22,
        'body_font': '宋体',
        'body_size': 12,
        'line_spacing': 1.5,
        'image_width': 5.91,
        'image_height': 4.43
    },
    'official': {
        'heading_font': '黑体',
        'heading_size': 22,
        'body_font': '仿宋',
        'body_size': 14,
        'line_spacing': 1.5,
        'image_width': 5.91,
        'image_height': 4.43
    },
    'academic': {
        'heading_font': '黑体',
        'heading_size': 18,
        'body_font': '宋体',
        'body_size': 11,
        'line_spacing': 2.0,
        'image_width': 5.91,
        'image_height': 4.43
    },
    'business': {
        'heading_font': '黑体',
        'heading_size': 22,
        'body_font': '宋体',
        'body_size': 12,
        'line_spacing': 1.5,
        'image_width': 5.91,
        'image_height': 4.43
    },
    'casual': {
        'heading_font': '微软雅黑',
        'heading_size': 16,
        'body_font': '微软雅黑',
        'body_size': 11,
        'line_spacing': 1.15,
        'image_width': 6.0,
        'image_height': 5.0
    }
}

class AIWordFormatter:
    def __init__(self, input_file_path, use_ai=True, tone='no_preference', style_config=None, log_callback=None):
        """初始化 Word 格式化器，创建可控临时目录（避免权限问题）"""
        self.input_file = input_file_path
        self.doc = Document(input_file_path)
        self.use_ai = use_ai
        self.tone = tone
        self.ai_processor = AITextProcessor(tone=tone, log_callback=log_callback) if use_ai else None

        # 应用样式配置
        self.style_config = self._validate_style_config(style_config)

        # 统一图片尺寸（从配置中读取）
        self.image_width = Inches(self.style_config['image_width'])
        self.image_height = Inches(self.style_config['image_height'])
        
        # 可控临时目录（创建在输入文件同目录下，避免系统权限问题）
        self.temp_dir = Path(os.path.dirname(input_file_path)) / "docx_temp_images"
        self.temp_dir.mkdir(exist_ok=True, mode=0o777)  # 赋予最高权限
        
        logger.info(f"AIWordFormatter 初始化完成，文件: {input_file_path}, AI 启用: {use_ai}")
        logger.info(f"临时目录已创建：{self.temp_dir}")

    def analyze_document(self):
        """分析文档并返回统计信息"""
        word_count = 0
        paragraph_count = 0
        image_count = 0
        table_count = 0
        
        # 统计段落和字数
        for para in self.doc.paragraphs:
            paragraph_count += 1
            word_count += len(para.text.strip().split())
        
        # 统计表格
        table_count = len(self.doc.tables)
        
        # 统计图片
        for para in self.doc.paragraphs:
            para_xml = para._element.xml
            has_image = (
                '<w:drawing>' in para_xml or
                '<pic:pic>' in para_xml or
                'a:pic' in para_xml or
                '<v:shape' in para_xml or
                '<v:image' in para_xml or
                '<w:pict>' in para_xml
            )
            
            if has_image:
                image_count += 1

        # 也检查inline_shapes中的图片
        image_count += len(self.doc.inline_shapes)

        # 返回去重后的图片数量
        # 由于段落中的图片和inline_shapes可能重复计算，这里需要去重
        unique_images = set()
        for para in self.doc.paragraphs:
            para_xml = para._element.xml
            if '<w:drawing>' in para_xml or '<pic:pic>' in para_xml:
                unique_images.add(str(para_xml))
        
        image_count = len(unique_images) + len(self.doc.inline_shapes)

        return {
            'word_count': word_count,
            'paragraph_count': paragraph_count,
            'image_count': image_count,
            'table_count': table_count
        }

    def __del__(self):
        """析构函数：自动清理临时目录（程序结束时执行）"""
        if self.temp_dir.exists():
            try:
                shutil.rmtree(self.temp_dir)
                logger.info(f"临时目录已自动清理：{self.temp_dir}")
            except Exception as e:
                logger.warning(f"清理临时目录失败: {e}，可手动删除：{self.temp_dir}")

    def _validate_style_config(self, config):
        """
        验证并清理样式配置，确保所有必需的键都存在且值有效
        :param config: 用户提供的样式配置字典
        :return: 验证后的完整样式配置字典
        """
        # 如果没有提供配置，使用默认模板
        if not config:
            logger.info("未提供样式配置，使用默认模板")
            return STYLE_TEMPLATES['default'].copy()

        # 如果是字符串，视为模板名称
        if isinstance(config, str):
            template_name = config
            if template_name in STYLE_TEMPLATES:
                logger.info(f"使用样式模板: {template_name}")
                return STYLE_TEMPLATES[template_name].copy()
            else:
                logger.warning(f"未知的样式模板: {template_name}，使用默认模板")
                return STYLE_TEMPLATES['default'].copy()

        # 如果是字典，合并默认值和用户提供的值
        if isinstance(config, dict):
            validated_config = STYLE_TEMPLATES['default'].copy()

            # 允许的字体选项
            valid_fonts = ['黑体', '宋体', '楷体', '仿宋', '微软雅黑']
            # 支持所有中文字号对应的磅值和中文名称
            valid_heading_sizes = list(CHINESE_FONT_SIZES.values()) + list(CHINESE_FONT_SIZES.keys())  # 包含数值和中文名称
            valid_body_sizes = list(CHINESE_FONT_SIZES.values()) + list(CHINESE_FONT_SIZES.keys())    # 包含数值和中文名称
            valid_line_spacings = [1.0, 1.15, 1.5, 2.0, 2.5, 3.0]
            # 预设图片尺寸选项（英寸）
            valid_image_widths = [4.0, 5.0, 5.91, 6.0, 7.0]
            valid_image_heights = [3.0, 4.0, 4.43, 5.0, 6.0]
            # 自定义图片尺寸的合理范围（英寸，1-10英寸 ≈ 2.54-25.4厘米）
            min_image_size = 1.0
            max_image_size = 10.0

            # 辅助函数：将中文字号代码或数值转换为磅值
            def convert_font_size(value):
                """转换字号：支持中文代码（'xiaosi'）或数值"""
                original_value = value  # 保存原始值用于日志

                if isinstance(value, str):
                    # 如果是中文字号代码（如 'xiaosi'）
                    if value in CHINESE_FONT_SIZES:
                        converted = CHINESE_FONT_SIZES[value]
                        logger.info(f"字号转换: 中文代码 '{value}' -> {converted} pt")
                        return converted
                    # 如果是纯数字字符串
                    try:
                        converted = float(value)
                        logger.info(f"字号转换: 数字字符串 '{value}' -> {converted} pt")
                        return converted
                    except ValueError:
                        logger.warning(f"字号转换失败: 无法识别的值 '{value}'")
                        return None
                elif isinstance(value, (int, float)):
                    logger.info(f"字号: 数值 {value} pt")
                    return float(value)

                logger.warning(f"字号转换失败: 值为空或类型错误")
                return None

            # 验证并更新每个键
            if 'heading_font' in config and config['heading_font'] in valid_fonts:
                validated_config['heading_font'] = config['heading_font']

            if 'heading_size' in config:
                original_size = config['heading_size']
                size = convert_font_size(original_size)
                if size is not None:
                    # Always store the numeric value for Pt() to use
                    if size in valid_heading_sizes or size in CHINESE_FONT_SIZES.values():
                        validated_config['heading_size'] = size  # Always store numeric value
                        logger.info(f"✅ 标题字号已应用: {original_size} -> {size} pt")
                    else:
                        logger.warning(f"⚠️ 标题字号无效或不在允许范围内: {original_size}")
                else:
                    logger.warning(f"⚠️ 标题字号无效或不在允许范围内: {original_size}")

            if 'body_font' in config and config['body_font'] in valid_fonts:
                validated_config['body_font'] = config['body_font']

            if 'body_size' in config:
                original_size = config['body_size']
                size = convert_font_size(original_size)
                if size is not None:
                    # Always store the numeric value for Pt() to use
                    if size in valid_body_sizes or size in CHINESE_FONT_SIZES.values():
                        validated_config['body_size'] = size  # Always store numeric value
                        logger.info(f"✅ 正文字号已应用: {original_size} -> {size} pt")
                    else:
                        logger.warning(f"⚠️ 正文字号无效或不在允许范围内: {original_size}")
                else:
                    logger.warning(f"⚠️ 正文字号无效或不在允许范围内: {original_size}")

            if 'line_spacing' in config and isinstance(config['line_spacing'], (int, float)):
                spacing = float(config['line_spacing'])
                if spacing in valid_line_spacings:
                    validated_config['line_spacing'] = spacing

            if 'image_width' in config and isinstance(config['image_width'], (int, float)):
                width = float(config['image_width'])
                # 接受预设值或自定义值（在合理范围内）
                if width in valid_image_widths or (min_image_size <= width <= max_image_size):
                    validated_config['image_width'] = width

            if 'image_height' in config and isinstance(config['image_height'], (int, float)):
                height = float(config['image_height'])
                # 接受预设值或自定义值（在合理范围内）
                if height in valid_image_heights or (min_image_size <= height <= max_image_size):
                    validated_config['image_height'] = height

            logger.info(f"样式配置已验证: {validated_config}")
            return validated_config

        # 无法识别的配置类型，使用默认值
        logger.warning(f"无法识别的配置类型: {type(config)}，使用默认模板")
        return STYLE_TEMPLATES['default'].copy()

    def _ensure_dir_writable(self, dir_path):
        """确保目录存在且可写（解决输出权限问题）"""
        dir_path = Path(dir_path)
        dir_path.mkdir(parents=True, exist_ok=True, mode=0o777)
        if not os.access(dir_path, os.W_OK):
            raise PermissionError(f"无写入权限：{dir_path}")
        return dir_path

    def _extract_images_from_docx(self):
        """从 docx 提取图片（手动写入，无权限问题，提升性能）"""
        image_paths = []
        try:
            with zipfile.ZipFile(self.input_file, 'r') as zip_file:
                # 遍历所有 media 文件夹下的图片
                for file_name in zip_file.namelist():
                    if file_name.startswith('word/media/') and not file_name.endswith('/'):
                        # 简化文件名：避免嵌套路径，提升读取速度
                        img_suffix = file_name.split('.')[-1]
                        img_filename = f"img_{len(image_paths)}.{img_suffix}"
                        img_path = self.temp_dir / img_filename
                        
                        # 手动写入文件（替代 zip.extract，避免权限问题，提升性能）
                        with open(img_path, 'wb') as f:
                            f.write(zip_file.read(file_name))
                        
                        image_paths.append(str(img_path))
                        logger.debug(f"提取图片: {file_name} -> {img_path}")

            logger.info(f"共提取到 {len(image_paths)} 张图片")
            return image_paths

        except Exception as e:
            logger.error(f"提取图片失败: {str(e)}")
            return []

    def _set_text_paragraph_style(self, para):
        """设置文本段落样式（标题/正文区分）"""
        para_text = para.text
        if para_text.startswith(("第", "一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、")):
            # 标题样式（从配置中读取）
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.name = self.style_config['heading_font']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.style_config['heading_font'])
                run.font.size = Pt(self.style_config['heading_size'])
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            # 正文样式（从配置中读取）
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para.paragraph_format.line_spacing = self.style_config['line_spacing']
            para.paragraph_format.first_line_indent = Pt(21.0)

            for run in para.runs:
                run.font.name = self.style_config['body_font']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.style_config['body_font'])
                run.font.size = Pt(self.style_config['body_size'])
                run.font.color.rgb = RGBColor(68, 68, 68)

    def _process_with_ai(self):
        """启用 AI 时：修复图片丢失 + 无重复日志 + 性能优化"""
        logger.info("启用 AI 处理模式")

        # 1. 提取所有图片（可控临时目录，无权限问题）
        image_paths = self._extract_images_from_docx()
        total_images = len(image_paths)
        
        # 2. 检测所有图片段落和文本段落
        original_paragraphs = list(self.doc.paragraphs)
        image_para_indices = []  # 保存图片段落的原始位置
        pure_texts = []
        text_para_indices = []   # 保存文本段落的原始位置

        for idx, para in enumerate(original_paragraphs):
            para_xml = para._element.xml
            para_text = para.text.strip()
            
            has_image = (
                '<w:drawing>' in para_xml or
                '<pic:pic>' in para_xml or
                '<v:shape' in para_xml or
                '<v:image' in para_xml or
                '<w:pict>' in para_xml
            )
            
            if has_image:
                image_para_indices.append(idx)
                logger.info(f"图片段落 {idx}：文本={para_text[:20]}...")
            elif para_text:
                pure_texts.append(para_text)
                text_para_indices.append(idx)

        logger.info(f"原始文档：{len(image_para_indices)} 个图片段落，{len(pure_texts)} 个纯文本段落")

        # 3. AI 处理文本（无重复日志 + 兜底逻辑）
        processed_text_blocks = []
        if pure_texts:
            merged_text = "\n".join(pure_texts)
            logger.info(f"AI 处理文本长度：{len(merged_text)} 字符")
            try:
                # 调用 AI 处理（AI 内部已打印日志，此处不重复打印）
                processed_text = self.ai_processor.process_text(merged_text)
                
                # 拆分文本块（保留原有逻辑，提升分段准确性）
                if not processed_text or processed_text.strip() == "":
                    logger.error("AI 返回空内容，使用原始文本")
                    processed_text_blocks = pure_texts
                else:
                    logger.info(f"AI 返回原文内容:\n{processed_text}")
                    processed_text_blocks = [p.strip() for p in processed_text.split("\n\n") if p.strip()]
            except Exception as e:
                logger.error(f"AI 处理出错: {str(e)}，使用原始文本")
                processed_text_blocks = pure_texts
        else:
            logger.warning("无纯文本段落")

        logger.info(f"AI 处理后文本块数量：{len(processed_text_blocks)}")

        # 4. 重建文档：按原始段落顺序处理（支持 AI 返回多段落）
        new_doc = Document()
        img_idx = 0
        text_idx = 0
        
        for original_para_idx, para in enumerate(original_paragraphs):
            # 重建图片段落
            if original_para_idx in image_para_indices and img_idx < total_images:
                img_para = new_doc.add_paragraph()
                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                img_para.paragraph_format.space_after = Pt(12)
                img_para.paragraph_format.space_before = Pt(12)
                
                try:
                    img_run = img_para.add_run()
                    img_run.add_picture(
                        image_paths[img_idx],
                        width=self.image_width,
                        height=self.image_height
                    )
                    img_idx += 1
                    logger.info(f"插入图片 {img_idx}/{total_images}：{image_paths[img_idx-1]}")
                except Exception as e:
                    logger.warning(f"插入图片 {img_idx} 失败: {e}")
                    img_para.add_run().text = "[图片]"
            
            # 重建文本段落（按顺序插入所有 AI 处理后的段落）
            elif original_para_idx in text_para_indices and text_idx < len(processed_text_blocks):
                text_para = new_doc.add_paragraph(processed_text_blocks[text_idx])
                self._set_text_paragraph_style(text_para)
                text_idx += 1
            else:
                # 空段落或跳过
                new_doc.add_paragraph()

        # 如果 AI 返回了额外的段落（超过原始数量），追加到文档末尾
        while text_idx < len(processed_text_blocks):
            text_para = new_doc.add_paragraph(processed_text_blocks[text_idx])
            self._set_text_paragraph_style(text_para)
            text_idx += 1

        # 替换原文档
        self.doc = new_doc
        logger.info(f"重建完成：{text_idx} 个文本段落 + {img_idx} 个图片段落")

    def _process_without_ai(self):
        """禁用 AI 时：仅统一样式，提升处理速度"""
        logger.info("禁用 AI 处理模式，仅统一样式")
        for para in self.doc.paragraphs:
            self._set_text_paragraph_style(para)

    def _process_all_paragraphs(self):
        """统一处理所有段落（根据 AI 启用状态分支）"""
        logger.info("开始处理所有段落（文本+图片）")
        if self.use_ai:
            self._process_with_ai()
        else:
            self._process_without_ai()

    def _process_tables(self):
        """统一表格样式，提升文档整洁度"""
        table_count = len(self.doc.tables)
        logger.info(f"处理 {table_count} 个表格")
        for i, table in enumerate(self.doc.tables):
            table.style = 'Table Grid'
            table.autofit = False
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(6.5)
                    for para in cell.paragraphs:
                        self._set_text_paragraph_style(para)

    def _process_images(self):
        """规范图片对齐和大小，提升文档美观度"""
        logger.info("开始处理文档中的图片")
        image_count = 0
        for shape in self.doc.inline_shapes:
            try:
                shape.width = self.image_width
                shape.height = self.image_height
                image_count += 1
            except Exception as e:
                logger.warning(f"设置图片大小失败: {str(e)}")

        for paragraph in self.doc.paragraphs:
            para_xml = paragraph._element.xml
            has_image = (
                '<w:drawing>' in para_xml or
                '<pic:pic>' in para_xml or
                'a:pic' in para_xml or
                '<v:shape' in para_xml or
                '<v:image' in para_xml or
                '<w:pict>' in para_xml
            )
            
            if has_image:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.paragraph_format.space_before = Pt(12)

        logger.info(f"图片处理完成，共规范化 {image_count} 张图片")

    def format(self, output_file_path):
        """主格式化逻辑：整合所有步骤，确保文件正常输出"""
        output_file_path = Path(output_file_path)
        logger.info(f"开始格式化文件，输入: {self.input_file}，输出: {output_file_path}")
        
        # 确保输出目录可写（解决权限问题）
        try:
            output_dir = self._ensure_dir_writable(output_file_path.parent)
        except PermissionError as e:
            logger.error(f"输出路径权限检查失败: {e}")
            # 降级方案：输出到桌面（确保可写）
            desktop_path = Path(os.path.expanduser("~")) / "Desktop"
            output_file_path = desktop_path / output_file_path.name
            logger.warning(f"降级到桌面输出：{output_file_path}")
            self._ensure_dir_writable(desktop_path)

        # 执行段落、表格、图片处理
        self._process_all_paragraphs()
        logger.info("段落处理完成")
        
        self._process_tables()
        logger.info("表格处理完成")
        
        self._process_images()
        logger.info("图片处理完成")
        
        # 保存文件（兜底重试，避免临时锁定）
        try:
            self.doc.save(str(output_file_path))
            logger.info(f"格式化完成，文件已保存: {output_file_path}")
            return str(output_file_path)
        except Exception as e:
            logger.error(f"保存文件失败: {e}")
            raise