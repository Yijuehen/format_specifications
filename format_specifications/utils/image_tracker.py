"""
Image tracking and reinsertion utilities for template mode.

This module provides functionality to extract images from Word documents,
track their positions relative to text content, and intelligently
reinsert them into AI-generated template-based documents.
"""
import os
import shutil
import zipfile
import logging
from typing import List, Dict, Tuple

logger = logging.getLogger(__name__)


class DocumentImageTracker:
    """
    Extract and track image positions in Word documents.

    This class handles extracting all images from a source document,
    capturing their original positions and surrounding text context
    for later semantic matching to template sections.
    """

    def __init__(self, docx_path: str):
        """
        Initialize the image tracker.

        Args:
            docx_path: Path to the source Word document
        """
        self.docx_path = docx_path
        self.temp_dir = None
        self.images = []  # List[ImageMetadata]

    def extract_images_with_context(self) -> List[Dict]:
        """
        Extract all images with surrounding text context.

        This method:
        1. Opens the document
        2. Creates a temporary directory for extracted images
        3. Extracts all images using zipfile
        4. Maps each image to its paragraph position
        5. Captures surrounding text context for semantic matching

        Returns:
            List of image metadata dictionaries:
            {
                'image_path': str,  # Path to extracted image file
                'paragraph_index': int,  # Original position in document
                'preceding_text': str,  # Text before image (200 chars)
                'following_text': str,  # Text after image (200 chars)
                'paragraph_text': str  # Full paragraph containing image
            }
        """
        from docx import Document

        logger.info(f"Starting image extraction from: {self.docx_path}")

        doc = Document(self.docx_path)
        total_paragraphs = len(doc.paragraphs)
        logger.info(f"Document has {total_paragraphs} paragraphs")

        # Create temp directory for images
        self.temp_dir = os.path.join(os.path.dirname(self.docx_path), 'docx_temp_images')
        os.makedirs(self.temp_dir, exist_ok=True)
        logger.info(f"Created temp directory: {self.temp_dir}")

        # Extract images using zipfile
        image_paths = self._extract_images_from_zipfile()
        logger.info(f"Extracted {len(image_paths)} image files from ZIP archive")

        if len(image_paths) == 0:
            logger.warning("No images found in document ZIP archive")
            return []

        # Map images to paragraphs
        image_index = 0
        paragraphs_with_images = 0

        for idx, para in enumerate(doc.paragraphs):
            if self._paragraph_has_image(para):
                paragraphs_with_images += 1
                logger.debug(f"Found image in paragraph {idx}: '{para.text[:50]}'")

                if image_index < len(image_paths):
                    self.images.append({
                        'image_path': image_paths[image_index],
                        'paragraph_index': idx,
                        'preceding_text': self._get_preceding_text(doc, idx),
                        'following_text': self._get_following_text(doc, idx),
                        'paragraph_text': para.text.strip()
                    })
                    image_index += 1
                else:
                    logger.warning(f"More image-containing paragraphs ({paragraphs_with_images}) than extracted images ({len(image_paths)})")

        logger.info(f"Image extraction complete:")
        logger.info(f"  - Total paragraphs scanned: {total_paragraphs}")
        logger.info(f"  - Paragraphs with images: {paragraphs_with_images}")
        logger.info(f"  - Images mapped to context: {len(self.images)}")
        logger.info(f"  - Images extracted from ZIP: {len(image_paths)}")

        # Warn if counts don't match
        if len(self.images) != len(image_paths):
            logger.warning(f"Image count mismatch: {len(self.images)} mapped vs {len(image_paths)} extracted. Some images may not be detected properly.")

        return self.images

    def _extract_images_from_zipfile(self) -> List[str]:
        """
        Extract images from docx file using zipfile.

        A .docx file is a ZIP archive containing media files in word/media/.

        Returns:
            List of paths to extracted image files
        """
        image_paths = []

        try:
            with zipfile.ZipFile(self.docx_path, 'r') as zip_ref:
                # Find all image files
                for file in zip_ref.namelist():
                    if file.startswith('word/media/'):
                        # Extract image
                        filename = os.path.basename(file)
                        extract_path = os.path.join(self.temp_dir, filename)

                        # Handle duplicate filenames
                        counter = 1
                        while os.path.exists(extract_path):
                            name, ext = os.path.splitext(filename)
                            extract_path = os.path.join(self.temp_dir, f"{name}_{counter}{ext}")
                            counter += 1

                        with zip_ref.open(file) as source:
                            with open(extract_path, 'wb') as target:
                                target.write(source.read())
                        image_paths.append(extract_path)

        except Exception as e:
            logger.error(f"Error extracting images: {e}")

        return image_paths

    def _paragraph_has_image(self, paragraph) -> bool:
        """
        Check if paragraph contains an image.

        Uses multiple detection strategies:
        1. XML pattern detection in paragraph element
        2. Run-level XML pattern detection
        3. Graphic object detection

        Args:
            paragraph: python-docx paragraph object

        Returns:
            True if paragraph contains image XML elements
        """
        try:
            # Method 1: Check paragraph-level XML for common image patterns
            para_xml = paragraph._element.xml
            xml_patterns = [
                '<w:drawing>',
                '<pic:pic>',
                '<v:shape',
                '<v:image',
                '<w:pict>',
                'a:graphic',  # Office 2007+ format
                'wp:inline',   # Inline floating image
                'wp:anchor'    # Anchored floating image
            ]
            if any(pattern in para_xml for pattern in xml_patterns):
                logger.debug(f"Found image in paragraph using XML pattern detection")
                return True

            # Method 2: Check runs within paragraph for images
            for run in paragraph.runs:
                run_xml = run._element.xml
                if any(pattern in run_xml for pattern in xml_patterns):
                    logger.debug(f"Found image in paragraph using run-level detection")
                    return True

                # Method 3: Check for graphic objects
                if 'graphic' in run_xml.lower() or 'blip' in run_xml.lower():
                    logger.debug(f"Found image in paragraph using graphic object detection")
                    return True

            return False

        except Exception as e:
            logger.warning(f"Error detecting image in paragraph: {e}")
            return False

    def _get_preceding_text(self, doc, current_idx: int, window: int = 3) -> str:
        """
        Get text from paragraphs before image.

        Args:
            doc: python-docx Document object
            current_idx: Current paragraph index
            window: Number of paragraphs to look back (default: 3)

        Returns:
            Combined text from preceding paragraphs
        """
        start_idx = max(0, current_idx - window)
        paragraphs = doc.paragraphs[start_idx:current_idx]
        return ' '.join([p.text.strip() for p in paragraphs if p.text.strip()])

    def _get_following_text(self, doc, current_idx: int, window: int = 3) -> str:
        """
        Get text from paragraphs after image.

        Args:
            doc: python-docx Document object
            current_idx: Current paragraph index
            window: Number of paragraphs to look ahead (default: 3)

        Returns:
            Combined text from following paragraphs
        """
        end_idx = min(len(doc.paragraphs), current_idx + window + 1)
        paragraphs = doc.paragraphs[current_idx + 1:end_idx]
        return ' '.join([p.text.strip() for p in paragraphs if p.text.strip()])

    def cleanup(self):
        """
        Remove temporary image files.

        Should be called in a finally block to ensure cleanup
        even if processing fails.
        """
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir)
                logger.debug(f"Cleaned up temp directory: {self.temp_dir}")
            except Exception as e:
                logger.warning(f"Failed to cleanup temp directory: {e}")


class ImageReinsertionStrategy:
    """
    Determine where to insert images in AI-generated content.

    This class implements semantic matching to place images in
    appropriate sections based on surrounding text context,
    section titles, and requirements.
    """

    @staticmethod
    def find_best_insertion_position(
        image_metadata: Dict,
        generated_sections: Dict[str, str],
        template
    ) -> Tuple[str, str]:
        """
        Find the best section and position to insert an image.

        Uses a multi-strategy approach:
        1. Primary: Keyword relevance scoring
        2. Fallback 1: First section with substantial content (>100 chars)
        3. Fallback 2: Last section
        4. Last resort: First section

        Args:
            image_metadata: Image context from original document
            generated_sections: AI-generated content by section ID
            template: Template object with section definitions

        Returns:
            Tuple of (section_id, position)
            - section_id: Which section to insert into (None if no match)
            - position: 'start' or 'end' (always 'end' for now)
        """
        best_section = None
        best_score = 0.0

        # Try to match based on keywords
        for section in template.sections:
            score = ImageReinsertionStrategy._calculate_relevance_score(
                image_metadata, section
            )
            if score > best_score:
                best_score = score
                best_section = section

        # If best match found and has meaningful score
        if best_section and best_score > 0:
            return (best_section.id, 'end')

        # Fallback 1: First section with substantial content
        for section_id, content in generated_sections.items():
            if len(content.strip()) > 100:
                return (section_id, 'end')

        # Fallback 2: Last section
        if generated_sections:
            last_section_id = list(generated_sections.keys())[-1]
            return (last_section_id, 'end')

        # Last resort: Any section
        if template.sections and len(template.sections) > 0:
            return (template.sections[0].id, 'end')

        return (None, 'end')

    @staticmethod
    def _calculate_relevance_score(image_metadata: Dict, section) -> float:
        """
        Calculate relevance score between image context and section.

        Scoring factors:
        1. Section title in preceding text (+0.5)
        2. Section title in following text (+0.5)
        3. Section requirement keywords (+0.1 each)

        Args:
            image_metadata: Image context dict with preceding/following text
            section: Template section object

        Returns:
            Relevance score (higher is better match)
        """
        score = 0.0

        # Check section title against preceding text
        if section.title.lower() in image_metadata['preceding_text'].lower():
            score += 0.5

        # Check section title against following text
        if section.title.lower() in image_metadata['following_text'].lower():
            score += 0.5

        # Check requirements
        if hasattr(section, 'requirements') and section.requirements:
            keywords = section.requirements.split()
            for keyword in keywords:
                if keyword.lower() in image_metadata['paragraph_text'].lower():
                    score += 0.1

        return score
