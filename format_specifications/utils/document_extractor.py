"""
Document text extraction utilities for template-based generation
"""
import logging
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Extract text content from Word documents for template-based generation
    """

    @staticmethod
    def extract_full_text(docx_path: str) -> str:
        """
        Extract all text from a Word document

        Args:
            docx_path: Path to the Word document

        Returns:
            Complete text content of the document
        """
        try:
            doc = Document(docx_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            return '\n\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting text from {docx_path}: {e}")
            raise

    @staticmethod
    def extract_by_headings(docx_path: str) -> Dict[str, str]:
        """
        Extract content organized by headings from a Word document

        Args:
            docx_path: Path to the Word document

        Returns:
            Dictionary mapping headings to their content
        """
        try:
            doc = Document(docx_path)
            sections = {}
            current_heading = None
            current_content = []

            for paragraph in doc.paragraphs:
                # Check if paragraph is a heading
                if paragraph.style.name.startswith('Heading'):
                    # Save previous section
                    if current_heading:
                        sections[current_heading] = '\n'.join(current_content).strip()

                    # Start new section
                    current_heading = paragraph.text.strip()
                    current_content = []
                else:
                    # Add to current section
                    text = paragraph.text.strip()
                    if text:
                        current_content.append(text)

            # Save last section
            if current_heading:
                sections[current_heading] = '\n'.join(current_content).strip()

            return sections

        except Exception as e:
            logger.error(f"Error extracting by headings from {docx_path}: {e}")
            raise

    @staticmethod
    def extract_with_structure(docx_path: str) -> List[Dict[str, any]]:
        """
        Extract document structure with paragraphs and metadata

        Args:
            docx_path: Path to the Word document

        Returns:
            List of dictionaries with text, style, and level information
        """
        try:
            doc = Document(docx_path)
            structure = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                structure.append({
                    'text': text,
                    'style': para.style.name,
                    'is_heading': para.style.name.startswith('Heading'),
                    'level': DocumentExtractor._get_heading_level(para.style.name)
                })

            return structure

        except Exception as e:
            logger.error(f"Error extracting structure from {docx_path}: {e}")
            raise

    @staticmethod
    def _get_heading_level(style_name: str) -> int:
        """
        Extract heading level from style name

        Args:
            style_name: Paragraph style name (e.g., 'Heading 1', 'Heading 2')

        Returns:
            Heading level (1-9) or 0 if not a heading
        """
        if style_name.startswith('Heading'):
            try:
                return int(style_name.split()[1])
            except (IndexError, ValueError):
                return 1
        return 0

    @staticmethod
    def extract_sections_by_keywords(docx_path: str, keywords: List[str]) -> Dict[str, str]:
        """
        Extract sections based on keyword matching in headings

        Args:
            docx_path: Path to the Word document
            keywords: List of keywords to match in headings

        Returns:
            Dictionary mapping matched keywords to their content
        """
        try:
            doc = Document(docx_path)
            sections = {keyword: [] for keyword in keywords}
            current_keyword = None

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()

                # Check if this heading matches any keyword
                if paragraph.style.name.startswith('Heading'):
                    matched_keyword = DocumentExtractor._match_keyword(text, keywords)
                    if matched_keyword:
                        current_keyword = matched_keyword
                        continue

                # Add content to current keyword section
                if current_keyword and text:
                    sections[current_keyword].append(text)

            # Convert lists to strings
            return {k: '\n'.join(v).strip() for k, v in sections.items()}

        except Exception as e:
            logger.error(f"Error extracting by keywords from {docx_path}: {e}")
            raise

    @staticmethod
    def _match_keyword(text: str, keywords: List[str]) -> Optional[str]:
        """
        Find if text contains any of the keywords

        Args:
            text: Text to search in
            keywords: List of keywords to match

        Returns:
            First matched keyword or None
        """
        text_lower = text.lower()
        for keyword in keywords:
            if keyword.lower() in text_lower:
                return keyword
        return None

    @staticmethod
    def get_document_statistics(docx_path: str) -> Dict[str, any]:
        """
        Get statistics about the document

        Args:
            docx_path: Path to the Word document

        Returns:
            Dictionary with document statistics
        """
        try:
            doc = Document(docx_path)
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]

            word_count = sum(len(p.text.split()) for p in paragraphs)
            char_count = sum(len(p.text) for p in paragraphs)

            heading_count = sum(1 for p in paragraphs if p.style.name.startswith('Heading'))

            return {
                'paragraph_count': len(paragraphs),
                'word_count': word_count,
                'char_count': char_count,
                'heading_count': heading_count
            }

        except Exception as e:
            logger.error(f"Error getting statistics for {docx_path}: {e}")
            raise
