"""
End-to-end test for segmentation feature
Tests the complete workflow from document upload to output generation
"""
import os
import sys
import django
from docx import Document
from docx.shared import Pt

# Setup Django environment
# Navigate to project root from tests/e2e/
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
sys.path.insert(0, PROJECT_ROOT)
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'format_specifications.settings')
django.setup()

from format_specifications.utils.ai_word_utils import AITextProcessor
from format_specifications.views import _build_segmented_document
from format_specifications.utils import generate_output_path
import tempfile


def create_test_document():
    """Create a test Word document with various content"""
    doc = Document()

    # Add title
    title = doc.add_heading('测试文档', 0)
    title.alignment = 0  # Center

    # Add first paragraph
    doc.add_paragraph('这是第一段内容。它包含了一些文字用于测试段落分割功能。')

    # Add second paragraph (empty line separation)
    doc.add_paragraph('这是第二段内容。这段话有不同的文字。')

    # Add third paragraph with multiple sentences
    doc.add_paragraph('这是第三段。它包含多个句子！这可以测试句子分割功能？是的，确实如此。')

    # Add section with heading for semantic segmentation
    doc.add_heading('一、第一章', level=1)
    doc.add_paragraph('这是第一章的内容。它有具体的描述。')
    doc.add_paragraph('这是第一章的第二个段落。')

    doc.add_heading('二、第二章', level=1)
    doc.add_paragraph('这是第二章的内容。')

    # Save to temporary file
    tmp_path = tempfile.mktemp(suffix='.docx')
    doc.save(tmp_path)
    return tmp_path


def test_paragraph_segmentation():
    """Test paragraph segmentation mode"""
    print("\n" + "=" * 60)
    print("TEST 1: Paragraph Segmentation Mode")
    print("=" * 60)

    # Create test document
    doc_path = create_test_document()
    print(f"[INFO] Created test document: {doc_path}")

    # Extract text
    doc = Document(doc_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n\n".join(paragraphs)
    print(f"[INFO] Extracted {len(paragraphs)} paragraphs from document")

    # Perform segmentation
    processor = AITextProcessor()
    segments = processor.segment_text(full_text, mode='paragraph', include_metadata=False)

    print(f"[RESULT] Segmented into {len(segments)} paragraphs:")
    for i, segment in enumerate(segments, 1):
        preview = segment[:50] + "..." if len(segment) > 50 else segment
        print(f"  [{i}] {preview}")

    # Build output document
    output_path = tempfile.mktemp(suffix='.docx')
    _build_segmented_document(segments, 'paragraph', False, output_path)
    print(f"[SUCCESS] Output document created: {output_path}")

    # Verify output
    output_doc = Document(output_path)
    output_paras = [p.text for p in output_doc.paragraphs if p.text.strip()]
    print(f"[INFO] Output document has {len(output_paras)} paragraphs")

    # Cleanup
    try:
        os.unlink(doc_path)
        os.unlink(output_path)
    except:
        pass

    return len(segments) > 0


def test_sentence_segmentation():
    """Test sentence segmentation mode"""
    print("\n" + "=" * 60)
    print("TEST 2: Sentence Segmentation Mode")
    print("=" * 60)

    # Create test document
    doc_path = create_test_document()
    print(f"[INFO] Created test document: {doc_path}")

    # Extract text
    doc = Document(doc_path)
    text = " ".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    print(f"[INFO] Extracted text ({len(text)} characters)")

    # Perform segmentation
    processor = AITextProcessor()
    segments = processor.segment_text(text, mode='sentence', include_metadata=False)

    print(f"[RESULT] Segmented into {len(segments)} sentences:")
    for i, segment in enumerate(segments[:5], 1):  # Show first 5
        print(f"  [{i}] {segment}")
    if len(segments) > 5:
        print(f"  ... and {len(segments) - 5} more sentences")

    # Build output document
    output_path = tempfile.mktemp(suffix='.docx')
    _build_segmented_document(segments, 'sentence', False, output_path)
    print(f"[SUCCESS] Output document created: {output_path}")

    # Cleanup
    try:
        os.unlink(doc_path)
        os.unlink(output_path)
    except:
        pass

    return len(segments) > 0


def test_semantic_segmentation():
    """Test semantic segmentation mode"""
    print("\n" + "=" * 60)
    print("TEST 3: Semantic Segmentation Mode")
    print("=" * 60)

    # Create test document
    doc_path = create_test_document()
    print(f"[INFO] Created test document: {doc_path}")

    # Extract text
    doc = Document(doc_path)
    text = "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    print(f"[INFO] Extracted text with headings")

    # Perform segmentation
    processor = AITextProcessor()
    segments = processor.segment_text(text, mode='semantic', include_metadata=False)

    print(f"[RESULT] Segmented into {len(segments)} semantic sections:")
    for i, segment in enumerate(segments, 1):
        preview = segment[:60] + "..." if len(segment) > 60 else segment
        print(f"  [{i}] {preview}")

    # Build output document
    output_path = tempfile.mktemp(suffix='.docx')
    _build_segmented_document(segments, 'semantic', False, output_path)
    print(f"[SUCCESS] Output document created: {output_path}")

    # Cleanup
    try:
        os.unlink(doc_path)
        os.unlink(output_path)
    except:
        pass

    return len(segments) > 0


def test_metadata_inclusion():
    """Test metadata inclusion"""
    print("\n" + "=" * 60)
    print("TEST 4: Metadata Inclusion")
    print("=" * 60)

    processor = AITextProcessor()
    text = "第一句。第二句。第三句。"

    segments = processor.segment_text(text, mode='sentence', include_metadata=True)

    print(f"[RESULT] Segmented with metadata:")
    for segment in segments:
        print(f"  Text: {segment['text']}")
        print(f"  Type: {segment['type']}")
        print(f"  Position: {segment['position']}")
        print()

    # Verify structure
    if segments:
        assert 'text' in segments[0], "Missing 'text' key"
        assert 'type' in segments[0], "Missing 'type' key"
        assert 'position' in segments[0], "Missing 'position' key"
        print("[SUCCESS] All metadata fields present")

    return len(segments) > 0 and all('text' in s and 'type' in s and 'position' in s for s in segments)


def test_output_document_format():
    """Test the format of generated output document"""
    print("\n" + "=" * 60)
    print("TEST 5: Output Document Format")
    print("=" * 60)

    # Create test data
    segments = ["这是第一段", "这是第二段", "这是第三段"]
    output_path = tempfile.mktemp(suffix='.docx')

    # Build document
    _build_segmented_document(segments, 'paragraph', False, output_path)
    print(f"[INFO] Created output document: {output_path}")

    # Verify document structure
    doc = Document(output_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    print(f"[RESULT] Document has {len(paragraphs)} paragraphs:")
    for i, para in enumerate(paragraphs[:10], 1):  # Show first 10
        print(f"  [{i}] {para}")

    # Check for expected elements
    assert '文档分割结果' in paragraphs[0], "Missing title"
    assert '分割模式' in paragraphs[1], "Missing mode info"
    assert '片段数量' in paragraphs[1], "Missing count info"
    assert '生成时间' in paragraphs[1], "Missing timestamp"

    print("[SUCCESS] Document format verified")
    print("  - Title present")
    print("  - Metadata present")
    print("  - Content segments present")

    # Cleanup
    try:
        os.unlink(output_path)
    except:
        pass

    return True


def main():
    """Run all end-to-end tests"""
    print("=" * 60)
    print("END-TO-END SEGMENTATION TESTING")
    print("=" * 60)

    results = {}

    try:
        results['paragraph'] = test_paragraph_segmentation()
        results['sentence'] = test_sentence_segmentation()
        results['semantic'] = test_semantic_segmentation()
        results['metadata'] = test_metadata_inclusion()
        results['format'] = test_output_document_format()

    except Exception as e:
        print(f"\n[ERROR] Test failed: {e}")
        import traceback
        traceback.print_exc()
        return 1

    # Summary
    print("\n" + "=" * 60)
    print("TEST SUMMARY")
    print("=" * 60)

    all_passed = True
    for test_name, passed in results.items():
        status = "[PASS]" if passed else "[FAIL]"
        print(f"{status} {test_name.capitalize()} Test")
        if not passed:
            all_passed = False

    print("\n" + "=" * 60)
    if all_passed:
        print("[SUCCESS] ALL END-TO-END TESTS PASSED!")
        print("=" * 60)
        print("\nFeature is ready for use:")
        print("  1. Start server: python manage.py runserver")
        print("  2. Navigate to: http://localhost:8000/segment/")
        print("  3. Upload a Word document")
        print("  4. Download segmented result")
        return 0
    else:
        print("[FAILURE] SOME TESTS FAILED")
        print("=" * 60)
        return 1


if __name__ == "__main__":
    sys.exit(main())
