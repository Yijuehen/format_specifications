"""
Unit tests for image tracker module
"""
import os
import sys
import tempfile
import shutil

# Setup Django environment
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
sys.path.insert(0, PROJECT_ROOT)
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'format_specifications.settings')

import django
django.setup()

from format_specifications.utils.image_tracker import DocumentImageTracker, ImageReinsertionStrategy
from docx import Document
from docx.shared import Inches


def test_image_tracker_initialization():
    """Test that DocumentImageTracker initializes correctly"""
    print("\n=== Test: Image Tracker Initialization ===")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name

    try:
        tracker = DocumentImageTracker(tmp_path)
        assert tracker.docx_path == tmp_path
        assert tracker.temp_dir is None
        assert tracker.images == []
        print("✓ DocumentImageTracker initialized correctly")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_image_extraction_from_empty_document():
    """Test extraction from document with no images"""
    print("\n=== Test: Extract from Empty Document ===")

    # Create a simple document with no images
    doc = Document()
    doc.add_paragraph("This is a test document")
    doc.add_paragraph("With no images")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name

    try:
        doc.save(tmp_path)

        tracker = DocumentImageTracker(tmp_path)
        images = tracker.extract_images_with_context()

        assert len(images) == 0, f"Expected 0 images, got {len(images)}"
        print(f"✓ Correctly extracted 0 images from document without images")

        tracker.cleanup()
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_paragraph_has_image_detection():
    """Test image detection in paragraphs"""
    print("\n=== Test: Paragraph Image Detection ===")

    tracker = DocumentImageTracker("dummy.docx")

    # Test with a simple paragraph (no image)
    doc = Document()
    para = doc.add_paragraph("Simple text paragraph")

    has_image = tracker._paragraph_has_image(para)
    assert not has_image, "Paragraph with only text should not have image"
    print("✓ Correctly detected paragraph without image")


def test_preceding_following_text_extraction():
    """Test extraction of surrounding text context"""
    print("\n=== Test: Surrounding Text Extraction ===")

    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.add_paragraph("Third paragraph")
    doc.add_paragraph("Fourth paragraph")
    doc.add_paragraph("Fifth paragraph")

    tracker = DocumentImageTracker("dummy.docx")

    # Test getting preceding text for paragraph at index 2
    preceding = tracker._get_preceding_text(doc, 2, window=2)
    assert "First paragraph" in preceding
    assert "Second paragraph" in preceding
    assert "Third paragraph" not in preceding
    print("✓ Correctly extracted preceding text")

    # Test getting following text for paragraph at index 2
    following = tracker._get_following_text(doc, 2, window=2)
    assert "Fourth paragraph" in following
    assert "Fifth paragraph" in following
    assert "Third paragraph" not in following
    print("✓ Correctly extracted following text")


def test_image_matching_strategy():
    """Test image to section matching logic"""
    print("\n=== Test: Image Matching Strategy ===")

    # Create mock image metadata
    image_meta = {
        'image_path': '/tmp/test.jpg',
        'paragraph_index': 5,
        'preceding_text': 'The sales chart shows growth',
        'following_text': 'as indicated in the performance data',
        'paragraph_text': 'Sales chart shows quarter over quarter growth'
    }

    # Create mock section
    class MockSection:
        def __init__(self, id, title, requirements=None):
            self.id = id
            self.title = title
            self.requirements = requirements

    section = MockSection('sales', '销售业绩', 'sales performance revenue')

    # Test relevance score calculation
    score = ImageReinsertionStrategy._calculate_relevance_score(image_meta, section)
    print(f"✓ Calculated relevance score: {score}")
    assert score > 0, "Should have positive score for matching keywords"


def test_cleanup():
    """Test temporary file cleanup"""
    print("\n=== Test: Temporary File Cleanup ===")

    # Create a temp directory manually
    temp_dir = tempfile.mkdtemp(prefix='test_cleanup_')

    # Create a mock image file
    test_file = os.path.join(temp_dir, 'test_image.jpg')
    with open(test_file, 'w') as f:
        f.write('test content')

    assert os.path.exists(temp_dir), "Temp directory should exist"
    assert os.path.exists(test_file), "Test file should exist"

    # Create a dummy tracker and manually set temp_dir
    tracker = DocumentImageTracker("dummy.docx")
    tracker.temp_dir = temp_dir

    # Cleanup
    tracker.cleanup()

    assert not os.path.exists(temp_dir), "Temp directory should be removed after cleanup"
    print("✓ Temporary files cleaned up correctly")


def run_all_tests():
    """Run all unit tests"""
    print("\n" + "="*60)
    print("Running Image Tracker Unit Tests")
    print("="*60)

    try:
        test_image_tracker_initialization()
        test_image_extraction_from_empty_document()
        test_paragraph_has_image_detection()
        test_preceding_following_text_extraction()
        test_image_matching_strategy()
        test_cleanup()

        print("\n" + "="*60)
        print("✅ All tests passed!")
        print("="*60)
        return True
    except AssertionError as e:
        print(f"\n❌ Test failed: {e}")
        return False
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = run_all_tests()
    sys.exit(0 if success else 1)
